"""Document processing utilities for SharePoint MCP server."""

import io
import os
import logging
from typing import Dict, Any, Optional

# Packages to support different file formats
try:
    import pandas as pd
    import docx
    from PyPDF2 import PdfReader
    import openpyxl  # noqa: F401
    import xlrd      # noqa: F401 — required for legacy .xls files

    HAS_DOCUMENT_LIBRARIES = True
except ImportError:
    HAS_DOCUMENT_LIBRARIES = False

try:
    import pyxlsb  # noqa: F401 — required for Binary Excel .xlsb files
    HAS_XLSB = True
except ImportError:
    HAS_XLSB = False

# Setup logging
logger = logging.getLogger("document_processor")


class DocumentProcessor:
    """Processor for various document types."""

    @staticmethod
    def check_dependencies():
        """Check if all required dependencies are installed."""
        if not HAS_DOCUMENT_LIBRARIES:
            logger.warning("Document processing libraries are not installed.")
            logger.warning(
                "Please install with: pip install pandas python-docx PyPDF2 openpyxl"
            )
            return False
        return True

    @staticmethod
    def process_document(
        content: bytes | None,
        filename: str,
        start_page: int = 1,
        end_page: int = None,
        file_path: str | None = None,
    ) -> Dict[str, Any]:
        """Process document content based on file type.

        Args:
            content:   Document content as bytes (small files loaded in RAM).
            filename:  Name of the file — used for format detection.
            start_page: First page/paragraph/row/line to extract (1-indexed).
            end_page:  Last to extract (inclusive). None = auto limit.
            file_path: Path to a temp file on disk (large files streamed from Graph).
                       When provided, 'content' must be None. The caller is
                       responsible for deleting the temp file.

        Returns:
            Processed document information dict.
        """
        if not DocumentProcessor.check_dependencies():
            return {"error": "Document processing libraries not installed"}

        file_ext = filename.lower().split(".")[-1] if "." in filename else ""

        # Determine the data source
        source_path: str | None = file_path        # disk path for large files
        source_bytes: bytes | None = content       # in-memory bytes for small files

        try:
            if file_ext == "csv":
                data = source_bytes or open(source_path, "rb").read()
                return DocumentProcessor._process_csv(
                    data, start_row=start_page, end_row=end_page
                )
            elif file_ext in ("xlsx", "xls"):
                return DocumentProcessor._process_excel(
                    source_bytes, file_ext=file_ext,
                    start_row=start_page, end_row=end_page,
                    file_path=source_path,
                )
            elif file_ext == "xlsb":
                return DocumentProcessor._process_xlsb(
                    source_bytes, start_row=start_page, end_row=end_page,
                    file_path=source_path,
                )
            elif file_ext == "docx":
                data = source_bytes or open(source_path, "rb").read()
                return DocumentProcessor._process_word(
                    data, start_para=start_page, end_para=end_page
                )
            elif file_ext == "pdf":
                data = source_bytes or open(source_path, "rb").read()
                return DocumentProcessor._process_pdf(
                    data, start_page=start_page, end_page=end_page
                )
            elif file_ext in ("txt", "md", "html", "htm"):
                data = source_bytes or open(source_path, "rb").read()
                return DocumentProcessor._process_text(
                    data, start_line=start_page, end_line=end_page
                )
            else:
                return {"error": f"Unsupported file type: {file_ext}"}
        except Exception as e:
            logger.error(f"Error processing document: {str(e)}")
            return {"error": str(e)}
        finally:
            # Always clean up temp file if one was provided
            if source_path:
                try:
                    os.remove(source_path)
                    logger.debug(f"Deleted temp file: {source_path}")
                except OSError:
                    pass

    @staticmethod
    def query_dataframe(
        file_path: str,
        filename: str,
        sheet_name: str | int = 0,
        query: str = "",
    ) -> Dict[str, Any]:
        """Execute a pandas expression on a dataset securely on the server.

        Args:
            file_path: Path to the temp file on disk.
            filename: Original filename to determine format.
            sheet_name: Sheet to load (for Excel files).
            query: Python expression to evaluate (e.g. "df['Sales'].sum()").
        """
        if not DocumentProcessor.check_dependencies():
            return {"error": "Document processing libraries not installed"}

        file_ext = filename.lower().split(".")[-1] if "." in filename else ""
        
        try:
            logger.info(f"Loading {filename} (sheet={sheet_name}) for query...")
            if file_ext == "csv":
                df = pd.read_csv(file_path)
            elif file_ext == "xlsb":
                df = pd.read_excel(file_path, sheet_name=sheet_name, engine="pyxlsb")
            elif file_ext == "xls":
                df = pd.read_excel(file_path, sheet_name=sheet_name, engine="xlrd")
            elif file_ext == "xlsx":
                df = pd.read_excel(file_path, sheet_name=sheet_name, engine="openpyxl")
            else:
                return {"error": f"Unsupported format for querying: {file_ext}"}
                
            # Execute the query securely
            logger.info(f"Executing query on DataFrame: {query}")
            
            # Simple local environment containing pandas and the dataframe
            local_env = {"df": df, "pd": pd}
            
            import numpy as np
            local_env["np"] = np

            result = eval(query, {"__builtins__": {}}, local_env)
            
            # Convert result to string or JSON serializable format
            if isinstance(result, pd.Series) or isinstance(result, pd.DataFrame):
                # For safety, limit large returns even from queries
                if len(result) > 1000:
                    return {"error": f"Query returned {len(result)} rows. Please aggregate further (e.g. use .head() or .sum())."}
                result_data = result.to_dict()
            else:
                # E.g. a float, int, str from aggregations
                result_data = str(result)
                
            return {
                "query": query,
                "result": result_data,
                "dataset_rows": len(df),
                "dataset_columns": list(df.columns)
            }
            
        except Exception as e:
            logger.error(f"Error querying dataframe: {str(e)}")
            return {"error": f"Failed to execute query: {str(e)}"}


    @staticmethod
    def _process_csv(
        content: bytes, start_row: int = 1, end_row: int = None
    ) -> Dict[str, Any]:
        """Process CSV content.

        Args:
            content: CSV file content
            start_row: First data row to return (1-indexed, excludes header). Default 1.
            end_row: Last data row to return (inclusive). None = start_row + 199 (200 rows).

        Returns:
            Processed data and analysis
        """
        df = pd.read_csv(io.BytesIO(content))
        total_rows = len(df)

        # Resolve row range (convert 1-indexed to 0-indexed)
        start_idx = max(0, start_row - 1)
        if end_row is None:
            end_idx = start_idx + 500  # default window: 500 rows
        else:
            end_idx = end_row
        end_idx = min(end_idx, total_rows)

        if start_idx >= total_rows:
            return {
                "error": f"start_row ({start_row}) exceeds total rows ({total_rows})",
                "total_rows": total_rows,
                "columns": list(df.columns),
            }

        df_slice = df.iloc[start_idx:end_idx]

        return {
            "type": "csv",
            "total_rows": total_rows,
            "extracted_range": {"start_row": start_idx + 1, "end_row": end_idx},
            "rows_returned": len(df_slice),
            "columns": list(df.columns),
            "data": df_slice.to_dict(orient="records"),
            "summary": {
                "numeric_columns": df.select_dtypes(include=["number"]).columns.tolist(),
                "missing_values": df.isnull().sum().to_dict(),
                "data_types": {col: str(dtype) for col, dtype in df.dtypes.items()},
            },
        }

    @staticmethod
    def _process_excel(
        content: bytes | None,
        file_ext: str = "xlsx",
        start_row: int = 1,
        end_row: int = None,
        file_path: str | None = None,
    ) -> Dict[str, Any]:
        """Process Excel content (.xlsx and legacy .xls).

        Args:
            content:  Excel file bytes (None when file_path is provided).
            file_ext: 'xlsx' or 'xls'.
            start_row: First data row to return per sheet (1-indexed). Default 1.
            end_row:  Last data row to return per sheet (inclusive).
                      None = start_row + 499. Pass total rows to read all.
            file_path: Path to temp file on disk (large files). Mutually
                       exclusive with content.
        """
        engine = "xlrd" if file_ext == "xls" else "openpyxl"
        logger.info(f"Reading Excel file with engine: {engine}")

        # Use file path directly if available — pandas reads from disk, not RAM
        source = file_path if file_path else io.BytesIO(content)

        sheets = {}
        for sheet_name, df in df_dict.items():
            total_rows = len(df)

            # Resolve row range
            start_idx = max(0, start_row - 1)
            if end_row is None:
                end_idx = start_idx + 500  # default window: 500 rows
            else:
                end_idx = end_row
            end_idx = min(end_idx, total_rows)

            df_slice = df.iloc[start_idx:end_idx] if start_idx < total_rows else df.iloc[0:0]

            # Sanitize column names (some legacy .xls files have unnamed columns)
            df.columns = [
                str(c) if not str(c).startswith("Unnamed") else f"Column_{i}"
                for i, c in enumerate(df.columns)
            ]
            df_slice.columns = df.columns

            sheets[sheet_name] = {
                "total_rows": total_rows,
                "extracted_range": {"start_row": start_idx + 1, "end_row": end_idx},
                "rows_returned": len(df_slice),
                "columns": list(df.columns),
                "data": df_slice.to_dict(orient="records"),
                "summary": {
                    "numeric_columns": df.select_dtypes(
                        include=["number"]
                    ).columns.tolist(),
                    "missing_values": df.isnull().sum().to_dict(),
                    "data_types": {col: str(dtype) for col, dtype in df.dtypes.items()},
                },
            }

        return {
            "type": "excel",
            "format": file_ext,
            "engine_used": engine,
            "sheet_count": len(sheets),
            "sheet_names": list(sheets.keys()),
            "sheets": sheets,
        }

    @staticmethod
    def _process_xlsb(
        content: bytes | None,
        start_row: int = 1,
        end_row: int = None,
        file_path: str | None = None,
    ) -> Dict[str, Any]:
        """Process Binary Excel content (.xlsb).

        Args:
            content:  Binary Excel bytes (None when file_path is provided).
            start_row: First data row per sheet (1-indexed). Default 1.
            end_row:  Last data row per sheet (inclusive).
                      None = start_row + 499. Pass total rows to read all.
            file_path: Path to temp file on disk (large files).
        """
        if not HAS_XLSB:
            return {
                "error": (
                    ".xlsb files require the pyxlsb library. "
                    "Install with: pip install pyxlsb"
                )
            }

        logger.info("Reading Binary Excel (.xlsb) file with pyxlsb engine")
        source = file_path if file_path else io.BytesIO(content)
        try:
            df_dict = pd.read_excel(source, sheet_name=None, engine="pyxlsb")
        except Exception as e:
            logger.error(f"Failed to read .xlsb file: {e}")
            return {"error": f"Failed to read .xlsb file: {e}"}

        sheets = {}
        for sheet_name, df in df_dict.items():
            total_rows = len(df)

            start_idx = max(0, start_row - 1)
            if end_row is None:
                end_idx = start_idx + 500
            else:
                end_idx = end_row
            end_idx = min(end_idx, total_rows)

            df_slice = df.iloc[start_idx:end_idx] if start_idx < total_rows else df.iloc[0:0]

            # Sanitize column names
            df.columns = [
                str(c) if not str(c).startswith("Unnamed") else f"Column_{i}"
                for i, c in enumerate(df.columns)
            ]
            df_slice.columns = df.columns

            sheets[sheet_name] = {
                "total_rows": total_rows,
                "extracted_range": {"start_row": start_idx + 1, "end_row": end_idx},
                "rows_returned": len(df_slice),
                "columns": list(df.columns),
                "data": df_slice.to_dict(orient="records"),
                "summary": {
                    "numeric_columns": df.select_dtypes(
                        include=["number"]
                    ).columns.tolist(),
                    "missing_values": df.isnull().sum().to_dict(),
                    "data_types": {col: str(dtype) for col, dtype in df.dtypes.items()},
                },
            }

        return {
            "type": "excel",
            "format": "xlsb",
            "engine_used": "pyxlsb",
            "sheet_count": len(sheets),
            "sheet_names": list(sheets.keys()),
            "sheets": sheets,
        }

    @staticmethod
    def _process_word(
        content: bytes, start_para: int = 1, end_para: int = None
    ) -> Dict[str, Any]:
        """Process Word document content.

        Args:
            content: Word document content
            start_para: First paragraph to return (1-indexed)
            end_para: Last paragraph to return (inclusive, None = start_para + 19)

        Returns:
            Processed text and structure
        """
        doc = docx.Document(io.BytesIO(content))

        all_paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        total_paragraphs = len(all_paragraphs)

        # Resolve range (convert 1-indexed to 0-indexed)
        start_idx = max(0, start_para - 1)
        if end_para is None:
            end_idx = start_idx + 20  # default window: 20 paragraphs
        else:
            end_idx = end_para  # end_para is 1-indexed inclusive → slice exclusive
        end_idx = min(end_idx, total_paragraphs)

        if start_idx >= total_paragraphs:
            return {
                "error": f"start_para ({start_para}) exceeds total paragraphs ({total_paragraphs})"
            }

        selected_paragraphs = all_paragraphs[start_idx:end_idx]

        tables = []
        for table in doc.tables:
            t_data = []
            for row in table.rows:
                t_data.append([cell.text for cell in row.cells])
            tables.append(t_data)

        # Extract document metadata
        core_properties = {}
        try:
            props = doc.core_properties
            core_properties = {
                "author": props.author or "",
                "title": props.title or "",
                "subject": props.subject or "",
                "created": str(props.created) if props.created else "",
                "modified": str(props.modified) if props.modified else "",
                "category": props.category or "",
            }
        except Exception as e:
            logger.warning(f"Error getting document properties: {str(e)}")

        # Extract headings/structure
        structure = []
        heading_styles = {"Heading 1", "Heading 2", "Heading 3", "Heading 4", "Heading 5", "Title"}
        for p in doc.paragraphs:
            if p.style.name in heading_styles and p.text.strip():
                level = 0 if p.style.name == "Title" else 1
                try:
                    level = int(p.style.name.split()[-1])
                except (ValueError, IndexError):
                    pass
                structure.append({"level": level, "text": p.text})

        return {
            "type": "word",
            "total_paragraph_count": total_paragraphs,
            "table_count": len(tables),
            "extracted_range": {"start_para": start_idx + 1, "end_para": end_idx},
            "content": selected_paragraphs,
            "tables": tables[:5],
            "properties": core_properties,
            "structure": structure,
        }

    @staticmethod
    def _process_pdf(
        content: bytes, start_page: int = 1, end_page: int = None
    ) -> Dict[str, Any]:
        """Process PDF content.

        Args:
            content: PDF file content
            start_page: First page to extract (1-indexed). Default: 1
            end_page: Last page to extract (inclusive, 1-indexed).
                      None = start_page + 9 (window of 10 pages).
                      Pass end_page equal to total pages to read to the end.

        Returns:
            Extracted text and metadata
        """
        pdf = PdfReader(io.BytesIO(content))
        total_pages = len(pdf.pages)

        # Resolve and clamp page range (convert 1-indexed to 0-indexed)
        start_idx = max(0, start_page - 1)
        if end_page is None:
            end_idx = start_idx + 10  # default window: 10 pages
        else:
            end_idx = end_page  # 1-indexed inclusive → slice exclusive
        end_idx = min(end_idx, total_pages)

        if start_idx >= total_pages:
            return {
                "error": (
                    f"start_page ({start_page}) exceeds total pages ({total_pages}). "
                    f"This PDF has {total_pages} pages."
                ),
                "total_pages": total_pages,
            }

        # Extract text for the requested range
        pages = []
        for i in range(start_idx, end_idx):
            page_text = pdf.pages[i].extract_text() or ""
            pages.append({"page": i + 1, "text": page_text})

        # Extract metadata
        metadata = {}
        if pdf.metadata:
            for key, value in pdf.metadata.items():
                if key.startswith("/"):
                    key = key[1:]
                if isinstance(value, (str, int, float, bool)) and key != "Trapped":
                    metadata[key] = str(value)

        # Extract form fields if present
        form_fields = []
        if hasattr(pdf, "get_fields"):
            fields = pdf.get_fields()
            if fields:
                for field_name, field_value in fields.items():
                    form_fields.append({"name": field_name, "value": str(field_value)})

        return {
            "type": "pdf",
            "total_pages": total_pages,
            "extracted_range": {"start_page": start_idx + 1, "end_page": end_idx},
            "pages_extracted": len(pages),
            "content": pages,
            "metadata": metadata,
            "form_fields": form_fields,
        }

    @staticmethod
    def _process_text(
        content: bytes, start_line: int = 1, end_line: int = None
    ) -> Dict[str, Any]:
        """Process text content.

        Args:
            content: Text file content
            start_line: First line to return (1-indexed). Default: 1
            end_line: Last line to return (inclusive). None = start_line + 29

        Returns:
            Processed text information
        """
        try:
            text = content.decode("utf-8")
        except UnicodeDecodeError:
            try:
                text = content.decode("latin-1")
            except Exception:
                return {"error": "Failed to decode text content"}

        lines = text.splitlines()
        total_lines = len(lines)

        # Resolve range
        start_idx = max(0, start_line - 1)
        if end_line is None:
            end_idx = start_idx + 30
        else:
            end_idx = end_line
        end_idx = min(end_idx, total_lines)

        if start_idx >= total_lines:
            return {
                "error": f"start_line ({start_line}) exceeds total lines ({total_lines})"
            }

        word_count = len(text.split())
        char_count = len(text)
        avg_line_length = char_count / total_lines if total_lines else 0

        is_markdown = text.count("#") > 0 and text.count("##") > 0
        is_html = text.count("<html") > 0 or text.count("<body") > 0

        return {
            "type": "text",
            "total_line_count": total_lines,
            "word_count": word_count,
            "character_count": char_count,
            "average_line_length": round(avg_line_length, 2),
            "extracted_range": {"start_line": start_idx + 1, "end_line": end_idx},
            "content": lines[start_idx:end_idx],
            "format": (
                "html" if is_html else ("markdown" if is_markdown else "plain_text")
            ),
        }
